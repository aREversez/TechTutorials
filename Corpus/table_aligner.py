import os
import docx
from docx.shared import Inches

# Prompt the user to enter the file paths
source_file_path = input("Enter the path to the source text file: ")
target_file_path = input("Enter the path to the target text file: ")

# Open the source and target docx files
source_doc = docx.Document(source_file_path)
target_doc = docx.Document(target_file_path)

# Extract text from the docx files
source_lines = [paragraph.text for paragraph in source_doc.paragraphs]
target_lines = [paragraph.text for paragraph in target_doc.paragraphs]

# Check if the number of lines in both files is the same
if len(source_lines) != len(target_lines):
    print("Warning: The number of lines in the source and target files is different.")

# Create a new docx document for the aligned texts
aligned_doc = docx.Document()

# Set up the table with two columns
table = aligned_doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Iterate over the lines and add the aligned sentences to the table
for source_line, target_line in zip(source_lines, target_lines):
    row_cells = table.add_row().cells
    row_cells[0].text = source_line.strip()
    row_cells[1].text = target_line.strip()

# Get the directory path of the source file
source_dir = os.path.dirname(source_file_path)

# Save the aligned docx file in the same directory as the source file
output_file_path = os.path.join(source_dir, 'table_aligned_texts.docx')
aligned_doc.save(output_file_path)

print(f"Aligned texts have been written to '{output_file_path}'")
